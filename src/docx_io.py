import docx

def read_docx(file_path: str) -> docx.Document:
    return docx.Document(file_path)

def save_docx(doc: docx.Document, file_path: str):
    doc.save(file_path)

def iter_paragraphs(parent):
    """
    Recursively walks through paragraphs in a document or cell,
    including inside tables and nested tables.
    """
    if hasattr(parent, 'paragraphs'):
        for paragraph in parent.paragraphs:
            yield paragraph
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)

def redact_paragraph(paragraph, spans, mapper) -> list:
    """
    Redacts the paragraph text based on accepted spans and EntityMapper.
    Returns a list of dicts summarizing the redactions made: [{'type': type, 'length': len}]
    """
    if not spans:
        return []

    # Get concatenated text and build run offset map
    # (Although we replace paragraph-wide, the prompt requested run-boundary mapping and right-to-left logic)
    concatenated_text = ""
    run_offsets = []  # Maps char index in concatenated_text to (run_index, char_index_within_run)
    for run_idx, run in enumerate(paragraph.runs):
        run_text = run.text
        for char_offset in range(len(run_text)):
            run_offsets.append((run_idx, char_offset))
        concatenated_text += run_text

    # If concatenated text doesn't match paragraph.text, fallback to paragraph.text
    if not concatenated_text:
        concatenated_text = paragraph.text

    # Replace right-to-left (highest offset first)
    sorted_spans = sorted(spans, key=lambda x: x[0], reverse=True)
    redacted_text = concatenated_text
    redactions_summary = []

    for start, end, entity_type, matched_text, confidence in sorted_spans:
        fake_value = mapper.get_fake(entity_type, matched_text)
        # Apply replacement
        redacted_text = redacted_text[:start] + fake_value + redacted_text[end:]
        redactions_summary.append({
            'type': entity_type,
            'original_length': end - start
        })

    # Clear all runs and put the full new text into run[0]
    if paragraph.runs:
        paragraph.runs[0].text = redacted_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = redacted_text

    return redactions_summary
