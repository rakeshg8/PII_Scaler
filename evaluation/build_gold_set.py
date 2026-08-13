import json
import os
import re
import docx

def iter_paragraphs(parent):
    if hasattr(parent, 'paragraphs'):
        for paragraph in parent.paragraphs:
            yield paragraph
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)

def make_span(text, target, entity_type):
    spans = []
    for match in re.finditer(re.escape(target), text):
        spans.append({
            "start": match.start(),
            "end": match.end(),
            "type": entity_type,
            "text": target
        })
    return spans

def main():
    docx_path = "data/input/RHP.docx"
    if not os.path.exists(docx_path):
        print(f"Error: Document not found at {docx_path}")
        return

    doc = docx.Document(docx_path)
    paragraphs = list(iter_paragraphs(doc))

    gold_pages = []

    # 1. Page 1: Promoters list (index 165 to 167)
    p1_text = paragraphs[165].text + "\n" + paragraphs[166].text + "\n" + paragraphs[167].text
    p1_spans = []
    for name in ["Kushal Subbayya Hegde", "Pushpa Kushal Hegde", "Rajesh Kushal Hegde", "Rohit Kushal Hegde", "Rakhi Girija Shetty"]:
        p1_spans.extend(make_span(p1_text, name, "PERSON"))
    for org in ["Dhaulagiri Family Trust", "Everest Family Trust", "Makalu Family Trust", "Broad Family Trust", "Annapurna Family Trust", "Kanchenjunga Family Trust", "Waterloo Industrial Park VI Private Limited"]:
        p1_spans.extend(make_span(p1_text, org, "COMPANY"))
    gold_pages.append({
        "page_id": "page_1_promoters",
        "text": p1_text,
        "spans": p1_spans
    })

    # 2. Page 2: General Information Contact (index 27 to 29)
    p2_text = paragraphs[27].text + "\n" + paragraphs[28].text + "\n" + paragraphs[29].text
    p2_spans = []
    p2_spans.extend(make_span(p2_text, "Sarthak Malvadkar", "PERSON"))
    p2_spans.extend(make_span(p2_text, "+ 91 20 45051234", "PHONE"))
    p2_spans.extend(make_span(p2_text, "cs.connect@kshinternational.com", "EMAIL"))
    # Address with a spaced PIN: 411 045 (tests recall failure on spacing)
    p2_spans.extend(make_span(p2_text, "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune – 411 045, Maharashtra, India", "ADDRESS"))
    gold_pages.append({
        "page_id": "page_2_contact",
        "text": p2_text,
        "spans": p2_spans
    })

    # 3. Page 3: Director details (index 4991 to 4994)
    p3_text = (
        "Name: " + paragraphs[4991].text + "\n" +
        "Designation: " + paragraphs[4992].text + "\n" +
        "DIN: " + paragraphs[4993].text + "\n" +
        "Address: " + paragraphs[4994].text
    )
    p3_spans = []
    p3_spans.extend(make_span(p3_text, "Kushal Subbayya Hegde", "PERSON"))
    p3_spans.extend(make_span(p3_text, "S. no. 245/ 104, Pushpakamal, Deccan Gymkhana Society, lane no. 3 Prabhat Road, opposite PYC basketball court, Deccan Gymkhana, Pune – 411 004 Maharashtra, India", "ADDRESS"))
    gold_pages.append({
        "page_id": "page_3_director_1",
        "text": p3_text,
        "spans": p3_spans
    })

    # 4. Page 4: Director details (index 4995 to 4998)
    p4_text = (
        "Name: " + paragraphs[4995].text + "\n" +
        "Designation: " + paragraphs[4996].text + "\n" +
        "DIN: " + paragraphs[4997].text + "\n" +
        "Address: " + paragraphs[4998].text
    )
    p4_spans = []
    p4_spans.extend(make_span(p4_text, "Rajesh Kushal Hegde", "PERSON"))
    p4_spans.extend(make_span(p4_text, "12 Buena Monte, NCL co-operative housing society, Panchvati, Pashan, Pune – 411 008, Maharashtra, India", "ADDRESS"))
    gold_pages.append({
        "page_id": "page_4_director_2",
        "text": p4_text,
        "spans": p4_spans
    })

    # 5. Page 5: Registrar details (index 799 to 804)
    p5_text = (
        paragraphs[799].text + "\n" +
        paragraphs[800].text + "\n" +
        paragraphs[801].text + "\n" +
        paragraphs[802].text + "\n" +
        paragraphs[803].text + "\n" +
        paragraphs[804].text
    )
    p5_spans = []
    p5_spans.extend(make_span(p5_text, "MUFG Intime India Private Limited", "COMPANY"))
    p5_spans.extend(make_span(p5_text, "Link Intime India Private Limited", "COMPANY"))
    p5_spans.extend(make_span(p5_text, "Shanti Gopalkrishnan", "PERSON"))
    p5_spans.extend(make_span(p5_text, "+91 81081 14949", "PHONE"))
    p5_spans.extend(make_span(p5_text, "kshinternational.ipo@in.mpms.mufg.com", "EMAIL"))
    p5_spans.extend(make_span(p5_text, "C-101, Embassy 247\n1st Floor, L B S Marg, Vikhroli (West) Mumbai 400083, (Maharashtra), India", "ADDRESS"))
    gold_pages.append({
        "page_id": "page_5_registrar",
        "text": p5_text,
        "spans": p5_spans
    })

    # 6. Page 6: Banker details (index 807 to 812)
    p6_text = (
        paragraphs[807].text + "\n" +
        paragraphs[808].text + "\n" +
        paragraphs[809].text + "\n" +
        paragraphs[810].text + "\n" +
        paragraphs[811].text + "\n" +
        paragraphs[812].text
    )
    p6_spans = []
    p6_spans.extend(make_span(p6_text, "HDFC Bank Limited", "COMPANY"))
    p6_spans.extend(make_span(p6_text, "Eric Bacha", "PERSON"))
    p6_spans.extend(make_span(p6_text, "Sachin Gawade", "PERSON"))
    p6_spans.extend(make_span(p6_text, "Pravin Teli", "PERSON"))
    p6_spans.extend(make_span(p6_text, "Siddharth Jadhav", "PERSON"))
    p6_spans.extend(make_span(p6_text, "Tushar Gavankar", "PERSON"))
    p6_spans.extend(make_span(p6_text, "+91 22 30752929", "PHONE"))
    p6_spans.extend(make_span(p6_text, "+91 22 30752928", "PHONE"))
    p6_spans.extend(make_span(p6_text, "+91 22 30752914", "PHONE"))
    p6_spans.extend(make_span(p6_text, "siddharth.jadhav@hdfcbank.com", "EMAIL"))
    p6_spans.extend(make_span(p6_text, "sachin.gawade@hdfcbank.com", "EMAIL"))
    p6_spans.extend(make_span(p6_text, "eric.bacha@hdfcbank.com", "EMAIL"))
    p6_spans.extend(make_span(p6_text, "tushar.gavankar@hdfcbank.com", "EMAIL"))
    p6_spans.extend(make_span(p6_text, "pravin.teli2@hdfcbank.com", "EMAIL"))
    p6_spans.extend(make_span(p6_text, "Lodha I Think Techno Campus, O-3 Level\nNext to Kanjurmarg Railway Station, Kanjurmarg (East) Mumbai – 400042, Maharashtra, India", "ADDRESS"))
    gold_pages.append({
        "page_id": "page_6_banker",
        "text": p6_text,
        "spans": p6_spans
    })

    # 7. Page 7: Statutory Auditor Consent (index 863 to 868)
    p7_text = paragraphs[863].text + "\n" + paragraphs[868].text
    p7_spans = []
    p7_spans.extend(make_span(p7_text, "Kirtane & Pandit, LLP", "COMPANY"))
    gold_pages.append({
        "page_id": "page_7_auditor",
        "text": p7_text,
        "spans": p7_spans
    })

    # 8. Page 8: Share Acquisition dates next to names (index 3340 to 3344)
    # Testing non-DOB dates next to names
    p8_text = (
        paragraphs[3340].text + "\n" +
        paragraphs[3341].text + "\n" +
        paragraphs[3342].text + "\n" +
        paragraphs[3343].text + "\n" +
        paragraphs[3344].text
    )
    p8_spans = []
    p8_spans.extend(make_span(p8_text, "Kushal Subbayya Hegde", "PERSON"))
    gold_pages.append({
        "page_id": "page_8_share_acq",
        "text": p8_text,
        "spans": p8_spans
    })

    # 9. Page 9: Negative Control - Financial Info Table 39
    p9_text = (
        paragraphs[3454].text + "\n" +
        paragraphs[3455].text + "\n" +
        paragraphs[3456].text + "\n" +
        paragraphs[3457].text + "\n" +
        paragraphs[3458].text + "\n" +
        paragraphs[3459].text
    )
    gold_pages.append({
        "page_id": "page_9_neg_financials",
        "text": p9_text,
        "spans": []
    })

    # 10. Page 10: Negative Control - Risk Factors 1 (index 668)
    p10_text = paragraphs[668].text
    gold_pages.append({
        "page_id": "page_10_neg_risk1",
        "text": p10_text,
        "spans": []
    })

    # 11. Page 11: Negative Control - Risk Factors 2 (index 671)
    p11_text = paragraphs[671].text
    gold_pages.append({
        "page_id": "page_11_neg_risk2",
        "text": p11_text,
        "spans": []
    })

    # 12. Page 12: Exemptions and Company Description (index 23 to 24)
    p12_text = paragraphs[23].text + "\n" + paragraphs[24].text
    p12_spans = []
    p12_spans.extend(make_span(p12_text, "Bhandary Metal Extrusion Private Limited", "COMPANY"))
    gold_pages.append({
        "page_id": "page_12_exemptions",
        "text": p12_text,
        "spans": p12_spans
    })

    # 13. Page 13: Unusual Company formats
    p13_text = (
        "Consent is received from Kanj & Co. LLP, practicing company secretaries.\n"
        "We also work with Lalit Muljibhai Sarvaiya & Co. and M/s. Lalit Muljibhai Sarvaiya & Associates."
    )
    p13_spans = []
    p13_spans.extend(make_span(p13_text, "Kanj & Co. LLP", "COMPANY"))
    p13_spans.extend(make_span(p13_text, "Lalit Muljibhai Sarvaiya & Co.", "COMPANY"))
    p13_spans.extend(make_span(p13_text, "Lalit Muljibhai Sarvaiya & Associates", "COMPANY"))
    gold_pages.append({
        "page_id": "page_13_unusual_companies",
        "text": p13_text,
        "spans": p13_spans
    })

    # 14. Page 14: Multi-line address split across table cells (tests table boundary failure)
    p14_text = (
        "Company Registered Office details split across cells:\n"
        "Cell 1: Registered Office: 11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed\n"
        "Cell 2: Pune – 411 501 Maharashtra, India"
    )
    p14_spans = []
    p14_spans.extend(make_span(p14_text, "11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed", "ADDRESS"))
    p14_spans.extend(make_span(p14_text, "Pune – 411 501 Maharashtra, India", "ADDRESS"))
    gold_pages.append({
        "page_id": "page_14_split_address",
        "text": p14_text,
        "spans": p14_spans
    })

    # 15. Page 15: Director casing and word order difference (generalization test)
    p15_text = (
        "The board meeting was attended by Hegde, Kushal Subbayya who signed the report.\n"
        "Mrs. Pushpa Hegde was also present. Rajesh K. Hegde sent his apologies."
    )
    p15_spans = []
    p15_spans.extend(make_span(p15_text, "Hegde, Kushal Subbayya", "PERSON"))
    p15_spans.extend(make_span(p15_text, "Pushpa Hegde", "PERSON"))
    p15_spans.extend(make_span(p15_text, "Rajesh K. Hegde", "PERSON"))
    gold_pages.append({
        "page_id": "page_15_casing_order",
        "text": p15_text,
        "spans": p15_spans
    })

    # 16. Page 16: Synthetic technical PII testing (SSNs, Credit Cards, IP addresses, DOB context)
    p16_text = (
        "Server administrator IP addresses: 192.168.1.100 and 2001:0db8:85a3:0000:0000:8a2e:0370:7334.\n"
        "Customer representative details: SSN: 999-12-3456. Card number: 4111 1111 1111 1111.\n"
        "Personal profile: Kushal Subbayya Hegde, date of birth: October 15, 1985. Born on: 12-04-1980."
    )
    p16_spans = []
    p16_spans.extend(make_span(p16_text, "192.168.1.100", "IP_ADDRESS"))
    p16_spans.extend(make_span(p16_text, "2001:0db8:85a3:0000:0000:8a2e:0370:7334", "IP_ADDRESS"))
    p16_spans.extend(make_span(p16_text, "999-12-3456", "SSN"))
    p16_spans.extend(make_span(p16_text, "4111 1111 1111 1111", "CREDIT_CARD"))
    p16_spans.extend(make_span(p16_text, "Kushal Subbayya Hegde", "PERSON"))
    p16_spans.extend(make_span(p16_text, "October 15, 1985", "DATE_OF_BIRTH"))
    p16_spans.extend(make_span(p16_text, "12-04-1980", "DATE_OF_BIRTH"))
    gold_pages.append({
        "page_id": "page_16_synthetic_pii",
        "text": p16_text,
        "spans": p16_spans
    })

    # Save to gold_annotations.json
    output_path = "evaluation/gold_annotations.json"
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(gold_pages, f, indent=2)

    print(f"Successfully generated new gold dataset with {len(gold_pages)} pages at {output_path}")

if __name__ == "__main__":
    main()
